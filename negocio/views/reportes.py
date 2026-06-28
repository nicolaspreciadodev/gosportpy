import io
import csv
import json
import logging
from datetime import timedelta
from django.utils import timezone
from django.db.models import Sum, Count, Q
from django.shortcuts import render
from django.http import HttpResponse, FileResponse
from django.views import View
from core.mixins import RoleRequiredMixin
from canchas.models import Cancha
from negocio.models import Reserva, Torneo
from django.contrib.auth.mixins import LoginRequiredMixin

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import openpyxl
from docx import Document

logger = logging.getLogger(__name__)

# ====================================================================
# MIME TYPES CONSTANTES
# ====================================================================
MIME_PDF = 'application/pdf'
MIME_XLSX = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
MIME_DOCX = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'


def build_download_response(buffer, filename, content_type):
    """Construye un FileResponse seguro para descarga de archivos binarios."""
    buffer.seek(0)
    response = FileResponse(buffer, as_attachment=True, filename=filename, content_type=content_type)
    # FileResponse handles Content-Disposition appropriately across browsers
    return response


# ====================================================================
# REPORTES CSV (DUEÑO)
# ====================================================================

class ReporteReservasView(RoleRequiredMixin, View):
    allowed_roles = ['DUEÑO']
    
    def get(self, request):
        canchas_dueño = Cancha.objects.filter(dueño=request.user)
        reservas = Reserva.objects.filter(cancha__in=canchas_dueño).order_by('-fecha', '-hora')

        estado_filter = request.GET.get('estado')
        if estado_filter:
            reservas = reservas.filter(estado=estado_filter)

        fecha_filter = request.GET.get('fecha')
        if fecha_filter:
            reservas = reservas.filter(fecha=fecha_filter)

        cancha_filter = request.GET.get('cancha_id')
        if cancha_filter:
            try:
                reservas = reservas.filter(cancha__id=int(cancha_filter))
            except ValueError:
                pass

        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="reporte_reservas.csv"'

        writer = csv.writer(response)
        writer.writerow(['ID', 'Cancha', 'Deportista', 'Fecha', 'Hora', 'Estado', 'Pagado'])

        for r in reservas:
            writer.writerow([r.id, r.cancha.nombre, r.usuario.email, r.fecha, r.hora, r.estado, 'Si' if r.pagado else 'No'])

        return response


class ReporteTorneosView(RoleRequiredMixin, View):
    allowed_roles = ['DUEÑO']
    
    def get(self, request):
        torneos = Torneo.objects.filter(organizador=request.user).order_by('-fecha_inicio')

        torneo_q = request.GET.get('torneo_q')
        if torneo_q:
            torneos = torneos.filter(nombre__icontains=torneo_q)

        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="reporte_torneos.csv"'

        writer = csv.writer(response)
        writer.writerow(['ID', 'Nombre', 'Deporte', 'Estado', 'Equipos Inscritos', 'Max Equipos'])

        for t in torneos:
            writer.writerow([t.id, t.nombre, t.deporte.nombre, t.estado, t.equipos.count(), t.max_equipos])

        return response


# ====================================================================
# PANEL ANALYTICS (DUEÑO)
# ====================================================================

class PanelAnalyticsView(RoleRequiredMixin, View):
    """Panel estadístico para el DUEÑO."""
    allowed_roles = ['DUEÑO']
    
    def get(self, request):
        ahora = timezone.now()
        hace_30_dias = ahora - timedelta(days=30)
        
        canchas_dueño = Cancha.objects.filter(dueño=request.user)
        
        reservas_ingresos = Reserva.objects.filter(
            cancha__in=canchas_dueño
        ).filter(
            Q(pagado=True) | Q(estado='COMPLETADA')
        )
        
        ingresos_totales = reservas_ingresos.aggregate(
            total=Sum('factura__total')
        )['total'] or 0.0

        reservas_mes = reservas_ingresos.filter(
            fecha__year=ahora.year, 
            fecha__month=ahora.month
        )
        ingresos_mes = reservas_mes.aggregate(
            total=Sum('factura__total')
        )['total'] or 0.0
        
        reservas_ultimos_30 = Reserva.objects.filter(
            cancha__in=canchas_dueño,
            fecha__gte=hace_30_dias.date(),
            estado__in=['PROGRAMADA', 'COMPLETADA']
        )
        total_reservas_30_dias = reservas_ultimos_30.count()
        
        reservas_por_dia = reservas_ultimos_30.values('fecha').annotate(
            cantidad=Count('id')
        ).order_by('fecha')
        
        dias_labels = [r['fecha'].strftime('%d %b') for r in reservas_por_dia]
        dias_data = [r['cantidad'] for r in reservas_por_dia]
        
        canchas_populares = Cancha.objects.filter(
            dueño=request.user
        ).annotate(
            num_reservas=Count('reservas')
        ).filter(num_reservas__gt=0).order_by('-num_reservas')
        
        canchas_labels = [c.nombre for c in canchas_populares]
        canchas_data = [c.num_reservas for c in canchas_populares]

        context = {
            'ingresos_totales': float(ingresos_totales),
            'ingresos_mes': float(ingresos_mes),
            'total_reservas_30_dias': total_reservas_30_dias,
            'dias_labels_json': json.dumps(dias_labels),
            'dias_data_json': json.dumps(dias_data),
            'canchas_labels_json': json.dumps(canchas_labels),
            'canchas_data_json': json.dumps(canchas_data),
        }
        
        return render(request, 'analytics.html', context)


# ====================================================================
# EXPORTACIONES MULTIFORMATO DE RESERVAS (TODOS LOS ROLES)
# ====================================================================

def get_reservas_por_rol(user):
    """Devuelve las reservas permitidas dependiendo del perfil."""
    if user.is_superuser or user.rol == 'ADMIN':
        return Reserva.objects.select_related('cancha', 'usuario', 'factura').all().order_by('-fecha', '-hora')
    elif user.rol == 'DUEÑO':
        return Reserva.objects.select_related('cancha', 'usuario', 'factura').filter(cancha__dueño=user).order_by('-fecha', '-hora')
    else:
        return Reserva.objects.select_related('cancha', 'usuario', 'factura').filter(usuario=user).order_by('-fecha', '-hora')


class ReporteReservasPdfView(LoginRequiredMixin, View):
    def get(self, request):
        try:
            reservas = get_reservas_por_rol(request.user)
            buffer = io.BytesIO()

            p = canvas.Canvas(buffer, pagesize=letter)
            p.setFont("Helvetica-Bold", 16)
            p.drawString(100, 750, f"Reporte de Reservas - GoSport2 ({request.user.rol})")
            
            p.setFont("Helvetica", 10)
            y = 710
            for idx, r in enumerate(reservas):
                texto = f"{idx+1}. {r.cancha.nombre} | {r.usuario.username} | {r.fecha} {r.hora} | {r.estado} | {'Pagado' if r.pagado else 'Pendiente'}"
                p.drawString(60, y, texto)
                y -= 18
                if y < 50:
                    p.showPage()
                    p.setFont("Helvetica", 10)
                    y = 750

            p.showPage()
            p.save()

            return build_download_response(buffer, 'reporte_reservas.pdf', MIME_PDF)
        except Exception as e:
            logger.error(f"Error generando PDF de reservas: {e}")
            return HttpResponse(f"Error generando reporte: {e}", status=500)


class ReporteReservasExcelView(LoginRequiredMixin, View):
    def get(self, request):
        try:
            reservas = get_reservas_por_rol(request.user)

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Reservas"
            
            ws.append(["ID", "Cancha", "Deportista", "Fecha", "Hora", "Estado", "Monto Total", "Pagado"])
            
            for r in reservas:
                monto = float(r.factura.total) if hasattr(r, 'factura') and r.factura else 0.0
                ws.append([
                    r.id, r.cancha.nombre, r.usuario.email, 
                    str(r.fecha), str(r.hora), r.estado, monto, 'Si' if r.pagado else 'No'
                ])
                
            buffer = io.BytesIO()
            wb.save(buffer)

            return build_download_response(buffer, 'reporte_reservas.xlsx', MIME_XLSX)
        except Exception as e:
            logger.error(f"Error generando Excel de reservas: {e}")
            return HttpResponse(f"Error generando reporte: {e}", status=500)


class ReporteReservasWordView(LoginRequiredMixin, View):
    def get(self, request):
        try:
            reservas = get_reservas_por_rol(request.user)

            doc = Document()
            doc.add_heading(f'Reporte de Reservas ({request.user.rol}) - GoSport2', 0)
            
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = 'Cancha'
            hdr_cells[2].text = 'Deportista'
            hdr_cells[3].text = 'Fecha'
            hdr_cells[4].text = 'Estado'
            
            for r in reservas:
                row_cells = table.add_row().cells
                row_cells[0].text = str(r.id)
                row_cells[1].text = r.cancha.nombre
                row_cells[2].text = r.usuario.get_full_name() or r.usuario.username
                row_cells[3].text = f"{r.fecha} {r.hora.strftime('%H:%M')}"
                row_cells[4].text = str(r.estado)
                
            buffer = io.BytesIO()
            doc.save(buffer)

            return build_download_response(buffer, 'reporte_reservas.docx', MIME_DOCX)
        except Exception as e:
            logger.error(f"Error generando Word de reservas: {e}")
            return HttpResponse(f"Error generando reporte: {e}", status=500)


# ====================================================================
# HISTORIAL GLOBAL DE RESERVAS (SOLO SUPERADMIN)
# ====================================================================

from django.core.paginator import Paginator
from usuarios.models import CustomUser as UserModel


class AdminHistorialReservasView(View):
    """Vista exclusiva del SuperAdmin: historial cronológico de reservas de TODA la plataforma.

    Permite filtrar por:
        - Usuario (búsqueda por username o email)
        - Cancha (búsqueda por nombre)
        - Estado de reserva
        - Rango de fechas (desde / hasta)

    Muestra por reserva: Usuario, Cancha/Establecimiento, Horario (inicio – fin),
    Estado y Método de pago (desde Factura si existe).
    """

    def dispatch(self, request, *args, **kwargs):
        if not request.user.is_authenticated or not request.user.is_superuser:
            from django.http import HttpResponseForbidden
            return HttpResponseForbidden("Acceso restringido al SuperAdmin.")
        return super().dispatch(request, *args, **kwargs)

    def get(self, request):
        reservas = Reserva.objects.select_related(
            'usuario', 'cancha', 'cancha__dueño', 'factura'
        ).order_by('-fecha', '-hora')

        # ── Filtros ────────────────────────────────────────────────────
        usuario_q = request.GET.get('usuario_q', '').strip()
        cancha_q = request.GET.get('cancha_q', '').strip()
        estado_q = request.GET.get('estado', '').strip()
        fecha_desde = request.GET.get('fecha_desde', '').strip()
        fecha_hasta = request.GET.get('fecha_hasta', '').strip()

        if usuario_q:
            reservas = reservas.filter(
                Q(usuario__username__icontains=usuario_q) |
                Q(usuario__email__icontains=usuario_q) |
                Q(usuario__first_name__icontains=usuario_q) |
                Q(usuario__last_name__icontains=usuario_q)
            )
        if cancha_q:
            reservas = reservas.filter(cancha__nombre__icontains=cancha_q)
        if estado_q:
            reservas = reservas.filter(estado=estado_q)
        if fecha_desde:
            reservas = reservas.filter(fecha__gte=fecha_desde)
        if fecha_hasta:
            reservas = reservas.filter(fecha__lte=fecha_hasta)

        # ── Paginación ─────────────────────────────────────────────────
        paginator = Paginator(reservas, 25)
        page_number = request.GET.get('page', 1)
        page_obj = paginator.get_page(page_number)

        context = {
            'page_obj': page_obj,
            'total_reservas': reservas.count(),
            # Filtros activos (para repoblar el form)
            'usuario_q': usuario_q,
            'cancha_q': cancha_q,
            'estado_activo': estado_q,
            'fecha_desde': fecha_desde,
            'fecha_hasta': fecha_hasta,
            'estados': Reserva.ESTADO_CHOICES,
        }
        return render(request, 'negocio/admin_historial_reservas.html', context)
